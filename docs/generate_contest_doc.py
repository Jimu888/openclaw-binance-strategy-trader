from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def add_p(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)


def add_bullets(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


TITLE = "项目说明｜让我们用策略来交易（OpenClaw × Binance Skills）"


def main():
    doc = Document()

    add_title(doc, TITLE)

    add_p(doc, "")

    add_h1(doc, "1. 项目概述")
    add_p(
        doc,
        "本项目是一个面向 OpenClaw 用户的可视化策略交易工具，聚焦于“策略 → 计算 → 预览 → 风险确认 → 通过 Binance Spot 提交挂单”的完整链路，用于在 Binance Spot 市场执行常见策略。"
    )
    add_p(
        doc,
        "交易执行基于 Binance 官方 Skills（spot），并在 UI 与文档中强调最小权限、IP 白名单、二次确认等安全围栏，降低 AI 代理交易风险。"
    )
    add_p(doc, "此外，本项目提供两种使用方式：")
    add_bullets(doc, [
        "可视化界面（UI）：适合快速配置与策略预览",
        "提示词（Prompt）在 Agent 内使用：适合自动化工作流与展示“Agent + Skills”的能力",
    ])

    add_p(doc, "本项目支持三类策略：")
    add_numbered(doc, [
        "金字塔建仓（分层限价买单）",
        "金字塔止盈（分层限价卖单）",
        "恒定混合策略（Constant Mix / 再平衡）",
    ])

    add_h1(doc, "2. 适用场景与价值")
    add_bullets(doc, [
        "适合希望用纪律化策略执行分层挂单的交易者（减少情绪化交易）",
        "可视化输入把复杂拆单计算变成低门槛操作",
        "让普通用户也能使用机构/量化团队常用的策略（分层建仓、分层止盈、恒定混合再平衡），不需要写代码也能上手",
        "通过 OpenClaw + Binance Skills 形成“策略工具 + Agent 工作流”的一体化体验",
    ])

    add_h1(doc, "3. 核心功能（用户视角）")

    add_h2(doc, "3.1 金字塔建仓（分层限价买单）")
    add_p(doc, "用户输入：")
    add_bullets(doc, [
        "总投入金额（Quote 计价：默认 USDT，可选 USDC / FDUSD）",
        "分层层数",
        "价格范围（最高建仓价 → 最低建仓价）",
    ])
    add_p(doc, "系统输出：")
    add_bullets(doc, [
        "每层限价、每层投入金额",
        "换算后的下单数量（带币种单位）",
        "前 3 层明细预览 + “查看全部”弹窗查看完整挂单表",
        "风险勾选 + 二次确认弹窗 → 批量提交限价买单",
    ])
    add_p(doc, "校验与防呆：")
    add_bullets(doc, [
        "若最高建仓价高于现价：提示错误（避免限价单立即成交）",
        "必须勾选“我已确认风险 / 我理解会下单”才允许执行",
    ])

    add_h2(doc, "3.2 金字塔止盈（分层限价卖单）")
    add_p(doc, "用户输入：")
    add_bullets(doc, [
        "卖出总数量（Base 计价）",
        "分层层数",
        "价格范围（最低止盈价 → 最高止盈价）",
    ])
    add_p(doc, "系统输出：")
    add_bullets(doc, [
        "每层限价、每层卖出数量",
        "前 3 层预览 + 查看全部弹窗",
        "风险勾选 + 二次确认弹窗 → 批量提交限价卖单",
    ])
    add_p(doc, "校验与防呆：")
    add_bullets(doc, [
        "若最低止盈价不高于现价：提示错误（避免限价单立即成交）",
        "卖出总数量不得超过余额（接入 spot skill 后可实时校验）",
    ])

    add_h2(doc, "3.3 恒定混合策略（Constant Mix / 再平衡）")
    add_p(doc, "策略目标：维持固定比例（如 BTC 50% / USDT 50%），当偏离超过阈值时执行再平衡：高了卖、低了买。")
    add_p(doc, "支持配置：")
    add_bullets(doc, [
        "目标比例、偏离阈值",
        "最大单次交易金额",
        "检查频率（每 X 小时 / 每天 / 每周 / 每月）",
    ])

    add_h1(doc, "4. 支持资产与交易对")
    add_bullets(doc, [
        "Base：BTC / ETH / BNB",
        "Quote：默认 USDT，可选 USDC / FDUSD",
        "UI 中可切换 Base/Quote，策略表格与单位自动同步更新",
    ])

    add_h1(doc, "5. Prompt/Agent 使用方式（含示例提示词）")
    add_p(doc, "除了 UI，本项目也可用提示词驱动 OpenClaw Agent 直接调用 Binance 官方 spot skill 完成查询与下单。")

    add_h2(doc, "示例 1：查询价格")
    add_p(doc, "提示词：用 spot skill 查询 BTCUSDT 当前价格，并返回最新价格与时间戳。")

    add_h2(doc, "示例 2：生成金字塔建仓计划（纯计算）")
    add_p(doc, "提示词：给我生成 BTCUSDT 的金字塔建仓计划：总投入 5000 USDT，5 层，区间 90000 到 70000。输出每层：限价、投入金额、数量，并给出二次确认用的汇总信息。")

    add_h2(doc, "示例 3：执行金字塔建仓（spot skill 下 LIMIT 买单）")
    add_p(doc, "提示词：我已确认风险。请通过 Binance spot skill 为 BTCUSDT 按以下计划提交 5 笔 GTC LIMIT BUY 挂单（逐条列出）。提交后返回每笔订单的 orderId 和状态。（并在提示词里粘贴计划表：price + quantity）")

    add_h2(doc, "示例 4：恒定混合策略建议（可选执行）")
    add_p(doc, "提示词：读取我现货账户 BTC 与 USDT 持仓（spot skill /api/v3/account），按 BTC 50%/USDT 50% 目标计算当前偏离；如果偏离超过 5%，给出建议买卖方向与数量，并询问我是否确认下单。")

    add_h1(doc, "6. 安全设计（参赛重点）")
    add_numbered(doc, [
        "最小权限建议：只开现货交易 +（可选）读取，禁止提现",
        "IP 白名单提示：强烈建议开启，降低密钥泄露风险",
        "二次确认：风险勾选 + 弹窗确认后才提交挂单",
        "密钥不进入前端：不写入 UI、.env、仓库；提供 SECURITY 指南",
    ])

    add_h1(doc, "7. 技术实现（简述）")
    add_bullets(doc, [
        "前端：React（策略输入、分层计算、预览表格、二次确认弹窗）",
        "交易能力：通过 OpenClaw 体系调用 Binance 官方 spot skill（下单/查询）",
        "文档：README + SECURITY + ENV.example（指导安全配置与运行）",
    ])

    add_h1(doc, "8. 开源地址")
    add_p(doc, "GitHub：https://github.com/Jimu888/openclaw-binance-strategy-trader")
    add_p(doc, "作者：几木 @0xjimumu")
    add_p(doc, "X：https://x.com/0xjimumu")
    add_p(doc, "币安广场：https://www.binance.com/zh-CN/square/profile/0xjimumu")

    out_path = "docs/项目说明_参赛资料.docx"
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
